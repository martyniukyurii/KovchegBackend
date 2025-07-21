import io
import re
from typing import Dict, List, Any, Optional
from docxtpl import DocxTemplate
from docx import Document
from tools.logger import Logger

class DocxTemplateService:
    def __init__(self):
        self.logger = Logger()
    
    def extract_variables_from_docx(self, file_content: bytes) -> List[str]:
        """Витягти всі змінні {{variable}} з .docx файлу"""
        try:
            # Створюємо Document з байтів
            doc = Document(io.BytesIO(file_content))
            
            # Збираємо весь текст з документа
            full_text = []
            
            # Текст з параграфів
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            # Текст з таблиць
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            # Текст з headers та footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        full_text.append(paragraph.text)
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        full_text.append(paragraph.text)
            
            # Об'єднуємо весь текст
            complete_text = '\n'.join(full_text)
            
            # Шукаємо всі змінні у форматі {{variable}}
            variables = re.findall(r'\{\{\s*(\w+)\s*\}\}', complete_text)
            
            # Повертаємо унікальні змінні
            unique_variables = list(set(variables))
            
            self.logger.info(f"Знайдено {len(unique_variables)} змінних у .docx файлі: {unique_variables}")
            
            return unique_variables
            
        except Exception as e:
            self.logger.error(f"Помилка витягування змінних з .docx: {str(e)}")
            return []
    
    def generate_document_from_template(self, template_content: bytes, variables: Dict[str, Any]) -> Optional[bytes]:
        """Згенерувати документ з .docx шаблона"""
        try:
            # Створюємо DocxTemplate з байтів
            doc_template = DocxTemplate(io.BytesIO(template_content))
            
            # Рендеримо документ зі змінними
            doc_template.render(variables)
            
            # Зберігаємо у BytesIO
            output = io.BytesIO()
            doc_template.save(output)
            output.seek(0)
            
            self.logger.info(f"Документ успішно згенеровано з {len(variables)} змінними")
            
            return output.getvalue()
            
        except Exception as e:
            self.logger.error(f"Помилка генерації документа з .docx шаблона: {str(e)}")
            return None
    
    def validate_variables(self, template_content: bytes, provided_variables: Dict[str, Any]) -> Dict[str, Any]:
        """Перевірити чи всі необхідні змінні надані"""
        try:
            # Витягуємо всі змінні з шаблона
            required_variables = self.extract_variables_from_docx(template_content)
            
            # Перевіряємо які змінні відсутні
            missing_variables = []
            for var in required_variables:
                if var not in provided_variables:
                    missing_variables.append(var)
            
            # Перевіряємо які змінні зайві
            extra_variables = []
            for var in provided_variables.keys():
                if var not in required_variables:
                    extra_variables.append(var)
            
            return {
                "required_variables": required_variables,
                "provided_variables": list(provided_variables.keys()),
                "missing_variables": missing_variables,
                "extra_variables": extra_variables,
                "is_valid": len(missing_variables) == 0
            }
            
        except Exception as e:
            self.logger.error(f"Помилка валідації змінних: {str(e)}")
            return {
                "required_variables": [],
                "provided_variables": [],
                "missing_variables": [],
                "extra_variables": [],
                "is_valid": False,
                "error": str(e)
            }
    
    def get_template_preview(self, template_content: bytes) -> Dict[str, Any]:
        """Отримати попередній перегляд шаблона з інфо про змінні"""
        try:
            # Витягуємо змінні
            variables = self.extract_variables_from_docx(template_content)
            
            # Читаємо базову інформацію про документ
            doc = Document(io.BytesIO(template_content))
            
            # Рахуємо статистику
            paragraphs_count = len(doc.paragraphs)
            tables_count = len(doc.tables)
            
            # Отримуємо перші кілька параграфів як превью
            preview_text = ""
            for i, paragraph in enumerate(doc.paragraphs[:3]):
                if paragraph.text.strip():
                    preview_text += paragraph.text[:100] + "...\n"
                if i >= 2:
                    break
            
            return {
                "variables": variables,
                "variables_count": len(variables),
                "paragraphs_count": paragraphs_count,
                "tables_count": tables_count,
                "preview_text": preview_text
            }
            
        except Exception as e:
            self.logger.error(f"Помилка отримання превью шаблона: {str(e)}")
            return {
                "variables": [],
                "variables_count": 0,
                "paragraphs_count": 0,
                "tables_count": 0,
                "preview_text": "",
                "error": str(e)
            }

# Глобальний екземпляр сервісу
docx_template_service = DocxTemplateService() 